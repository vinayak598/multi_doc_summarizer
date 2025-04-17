import spacy
import streamlit as st
from transformers import pipeline
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from fpdf import FPDF

# Load SpaCy model for tokenization
nlp = spacy.load("en_core_web_sm")

# Load models once (Cache for efficiency)
@st.cache_resource
def load_summarizer(model_name):
    return pipeline("summarization", model=model_name)

# Model options
model_names = {
    "DistilBART": "sshleifer/distilbart-cnn-12-6",
    "BART Large": "facebook/bart-large-cnn",
    "T5 Base": "t5-base"
}

# --- Custom CSS Styling ---
st.markdown("""
    <style>
        .main {
            background-color: #f5f7fa;
            padding: 2rem;
        }
        .title {
            font-size: 2.5rem;
            font-weight: bold;
            text-align: center;
            color: #0f4c75;
        }
        .subtitle {
            font-size: 1.2rem;
            text-align: center;
            margin-bottom: 2rem;
            color: #3282b8;
        }
        .stButton>button {
            background-color: #3282b8;
            color: white;
            font-weight: bold;
            border-radius: 10px;
            padding: 0.5rem 1rem;
        }
        .stTextArea textarea {
            border-radius: 10px;
        }
        .summary-box {
            background-color: #ffffff;
            padding: 1.5rem;
            border-radius: 12px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
            margin-top: 1rem;
        }
    </style>
""", unsafe_allow_html=True)

# --- Header Section ---
st.markdown('<div class="title">🧠 Multi-Document Summarizer</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Summarize your text using Transformer models</div>', unsafe_allow_html=True)

# --- Model Selection ---
model_selection = st.selectbox("Select Summarization Model", list(model_names.keys()))
summarizer = load_summarizer(model_names[model_selection])

# --- System Requirements ---
st.markdown("""
### ⚙️ Model System Requirements:
1. **DistilBART** – Lightweight, good for small docs
2. **BART Large** – Better summaries, needs more RAM
3. **T5 Base** – Great for structured tasks

💡 All models support ~512–1024 tokens at once (~1-3 pages of text)
""", unsafe_allow_html=True)

# --- Model Input Guide ---
st.markdown("""
<details>
<summary>📌 <b>Model Input Guide (click to expand)</b></summary>

<p>Select a model based on your system's available RAM:</p>

<p>Based on the selected model, here is the model's specification:</p>

| Model        | Max Input Tokens | Pages  | Recommended RAM  |
|--------------|------------------|--------|------------------|
| DistilBART   | ~1024 tokens     | 2–3    | 4GB – 8GB        |
| BART Large   | ~1024 tokens     | 2–3    | 16GB – 32GB      |
| T5 Base      | ~512–1024 tokens | 1–2    | 8GB – 16GB       |

💡 Files above 5MB will be chunked automatically.

</details>

""", unsafe_allow_html=True)

# --- Manual Text Input ---
manual_text = st.text_area("Enter Text to Summarize", height=200)

# --- File Upload ---
uploaded_file = st.file_uploader("Or upload a PDF, DOCX, or TXT file", type=["pdf", "docx", "txt"])

# --- Chunking Long Text ---
def chunk_text(text, max_tokens=1024):
    words = text.split()
    chunks = []
    for i in range(0, len(words), max_tokens):
        chunk = " ".join(words[i:i+max_tokens])
        chunks.append(chunk)
    return chunks

# --- Extract Keywords ---
def extract_keywords(text, top_n=5):
    from sklearn.feature_extraction.text import TfidfVectorizer
    vectorizer = TfidfVectorizer(stop_words='english', max_features=top_n)
    tfidf_matrix = vectorizer.fit_transform([text])
    words = vectorizer.get_feature_names_out()
    return words

# --- Text Extraction ---
def extract_text_from_file(file):
    file_extension = file.name.split('.')[-1].lower()
    if file_extension == "pdf":
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    elif file_extension == "docx":
        doc = DocxDocument(file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif file_extension == "txt":
        return file.getvalue().decode("utf-8")
    return None

# --- Summarize Button ---
if st.button("Summarize"):
    with st.spinner("Summarizing..."):
        input_text = ""
        if manual_text.strip():
            input_text = manual_text.strip()
        elif uploaded_file:
            input_text = extract_text_from_file(uploaded_file)
        else:
            st.warning("Please enter text or upload a file.")
        
        if input_text:
            # Extractive Summary
            chunks = chunk_text(input_text)
            extractive_summary = []
            for chunk in chunks:
                if len(chunk.split()) < 10:
                    continue
                doc = nlp(chunk)
                sentences = [sent.text for sent in doc.sents]
                extractive_summary.append(" ".join(sentences[:3]))  # Top 3 sentences
            extractive_summary = " ".join(extractive_summary)

            # Abstractive Summary
            abstractive_summary = summarizer(extractive_summary, max_length=300, min_length=80, do_sample=False)[0]['summary_text']

            # Keywords
            keywords = extract_keywords(abstractive_summary)
            st.markdown("<b>🔑 Extracted Keywords:</b> " + ", ".join(keywords))

            # Display Summary
            st.markdown('<div class="summary-box"><b>📄 Final Summary:</b><br><br>' + abstractive_summary + '</div>', unsafe_allow_html=True)

            # Save as TXT
            txt_download = BytesIO()
            txt_download.write(abstractive_summary.encode())
            txt_download.seek(0)
            st.download_button("📄 Download Summary as TXT", txt_download, file_name="summary.txt", mime="text/plain")

            # Save as DOCX
            doc = DocxDocument()
            doc.add_heading("Generated Summary", level=1)
            doc.add_paragraph(abstractive_summary)
            docx_download = BytesIO()
            doc.save(docx_download)
            docx_download.seek(0)
            st.download_button("📄 Download Summary as DOCX", docx_download, file_name="summary.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            # Save as PDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, abstractive_summary)
            pdf_output = pdf.output(dest='S').encode('latin1')
            pdf_download = BytesIO(pdf_output)
            st.download_button("📄 Download Summary as PDF", pdf_download, file_name="summary.pdf", mime="application/pdf")

            # --- Credits and License ---
st.markdown("""
### 📜 Credits and License
This project was developed by **[JVR VINAYAK]**. The code and its content are licensed under the **MIT License**.

For more information, visit [Your GitHub](https://github.com/vinayak598).

""", unsafe_allow_html=True)